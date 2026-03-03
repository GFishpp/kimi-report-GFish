from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta

# 创建文档
doc = Document()

# 设置标题
title = doc.add_heading('每日币圈新闻报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 日期范围
yesterday = (datetime.now() - timedelta(days=1)).strftime('%Y-%m-%d')
today = datetime.now().strftime('%Y-%m-%d')
date_para = doc.add_paragraph(f'报告日期：{yesterday} 10:00 - {today} 10:00')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加分隔线
doc.add_paragraph('_' * 50)

# 一、市场概览
doc.add_heading('一、市场概览', 1)

market_items = [
    ('1. 全网爆仓数据', '据CoinAnk数据显示，过去24小时加密货币市场全网合约爆仓3.5亿美元，其中多单爆仓1.94亿美元，空单爆仓1.56亿美元。BTC爆仓总金额1.6亿美元，ETH爆仓总金额7066.71万美元。'),
    ('2. 币安新交易对上线', '币安宣布于2026年3月3日16:00(UTC+8)开放BNB/USD、ETH/USD、SOL/USD现货交易，同时上线这些交易对的Spot Algo Orders交易机器人。'),
    ('3. CME Group扩展加密期货产品线', 'CME Group上线Cardano(ADA)、Chainlink(LINK)、Stellar(XLM)期货（含大合约与微型、现金结算），与BTC/ETH/SOL/XRP产品合计可覆盖逾75%加密市值。Nasdaq CME Crypto Index期货拟于2026年3月16日推出（待监管批准）。'),
]

for title_text, content in market_items:
    p = doc.add_paragraph()
    p.add_run(title_text).bold = True
    doc.add_paragraph(content, style='List Bullet')

# 二、监管动态
doc.add_heading('二、监管动态', 1)

reg_items = [
    ('1. 香港稳定币牌照', '香港财政司司长陈茂波表示，香港已实施法币稳定币发行人发牌制度，将在3月发出首批牌照。政府将于年内提交数字资产政策条例草案，就数字资产交易及托管等服务提供者订立发牌制度。'),
    ('2. 欧盟银行联盟推欧元稳定币', '由ING、UniCredit、BNP Paribas、CaixaBank、BBVA等12家欧盟银行组建的Qivalis正与多家加密交易所、做市商洽谈，计划今年下半年推出符合MiCA监管的欧元锚定稳定币。该稳定币将按1:1由至少40%银行存款及高质量欧元区主权债券作为储备。'),
    ('3. 俄罗斯加密监管法案', '俄罗斯央行与财政部起草《数字货币和数字权利法》草案，拟于7月1日生效。草案将便利国家机关获取加密钱包与交易数据，仅持牌交易所可组织加密交易，非合格投资者单笔投资上限或为30万卢布（约3300美元）。'),
    ('4. 越南加密货币交易合法化', '越南计划投资部提议自2026年7月1日起，在金融中心（岘港市区域中心、胡志明市国际金融中心）正式实施加密资产和加密货币交易。'),
]

for title_text, content in reg_items:
    p = doc.add_paragraph()
    p.add_run(title_text).bold = True
    doc.add_paragraph(content, style='List Bullet')

# 三、机构动态
doc.add_heading('三、机构动态', 1)

inst_items = [
    ('1. ProCap Financial增持比特币', 'ProCap Financial（纳斯达克代码：BRR）宣布新增购入450枚Bitcoin，使其持币总量增至5,457枚。过去10天内，公司以明显低于净资产价值（NAV）的价格回购了782,408股普通股。'),
    ('2. Meta计划重启稳定币业务', 'Meta Platforms计划于今年下半年重新进入稳定币领域，前提是能成功与第三方公司整合。Meta已向第三方公司发出产品征求建议书（RFP），Stripe可能会成为Meta稳定币试点的合作方。'),
    ('3. 21Shares推出SUI ETF', '数字资产管理公司21Shares推出21Shares Spot SUI ETF，并已在纳斯达克挂牌交易。此前21Shares已于去年12月推出2X SUI杠杆ETF。'),
    ('4. Coinbase与雅虎财经合作', 'Coinbase已向所有美国用户开放股票交易功能，并与雅虎财经达成合作，将在后者的股票及加密资产页面添加"在Coinbase交易"按钮。'),
    ('5. Gate Ventures战略投资', 'Gate.com旗下风投Gate Ventures宣布战略投资比特币原生金融平台Sats Terminal，后者聚焦为BTC持有者提供自托管的交易、借贷与收益服务。'),
]

for title_text, content in inst_items:
    p = doc.add_paragraph()
    p.add_run(title_text).bold = True
    doc.add_paragraph(content, style='List Bullet')

# 四、链上数据
doc.add_heading('四、链上数据', 1)

onchain_items = [
    ('1. 爆仓数据分析', '过去24小时全网爆仓3.5亿美元，BTC爆仓1.6亿美元，ETH爆仓7066万美元。多单爆仓占主导，显示市场多头情绪受挫。'),
    ('2. Venice Token上涨', 'Venice AI创始人Erik Voorhees发文称OpenClaw文档显示已将Venice设为推荐模型提供商。行情显示，Venice Token(VVV)过去24小时涨14.3%，过去7天涨70.4%。'),
]

for title_text, content in onchain_items:
    p = doc.add_paragraph()
    p.add_run(title_text).bold = True
    doc.add_paragraph(content, style='List Bullet')

# 五、市场分析
doc.add_heading('五、市场分析', 1)

analysis_items = [
    ('1. 3月关键事件展望', 'PANews加密日历显示，3月加密市场关键事件包括：FOMC利率决议及鲍威尔新闻发布会、《Clarity法案》表决、美国2月非农及CPI数据、香港首批稳定币牌照发放、SUI/HYPE等代币大额解锁、Metaplanet股东大会、FTX新一轮资金分配。'),
    ('2. 纳斯达克进军预测市场', '纳斯达克公司计划推出允许对主要股票指数进行"是或否"投注的期权合约，将在纳斯达克100指数和纳斯达克100微型指数上上市"二元期权"。'),
]

for title_text, content in analysis_items:
    p = doc.add_paragraph()
    p.add_run(title_text).bold = True
    doc.add_paragraph(content, style='List Bullet')

# 六、宏观经济
doc.add_heading('六、宏观经济', 1)

macro_items = [
    ('1. 欧洲央行降息预期降温', '欧元区货币市场交易员几乎已排除欧洲央行在2026年降息的可能性。交易员们现在认为欧洲央行12月前降息的可能性约为8%，而上周五的预测约为40%。'),
    ('2. 美国财长表态', '美国财长贝森特表示，财政部正在终止对Anthropic产品的所有使用，包括其Claude平台。'),
    ('3. 白银价格波动', '现货白银向下触及89美元/盎司，日内跌5.25%。'),
]

for title_text, content in macro_items:
    p = doc.add_paragraph()
    p.add_run(title_text).bold = True
    doc.add_paragraph(content, style='List Bullet')

# 七、个人交易笔记
doc.add_heading('七、个人交易笔记', 1)
doc.add_paragraph('（此处留空供个人记录）')

# 添加页脚
doc.add_paragraph('_' * 50)
footer = doc.add_paragraph(f'报告生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
source = doc.add_paragraph('数据来源：PANews')
source.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存文档
filename = f'/root/.openclaw/workspace/crypto_report_{today.replace("-", "")}.docx'
doc.save(filename)
print(f'文档已保存至：{filename}')
